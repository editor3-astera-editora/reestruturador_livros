import docx
from pathlib import Path
from src.utils.logger import logger
import re

def add_markdown_paragraph(doc, markdown_text):
    """
    Adiciona um parágrafo ao documento, interpretando Markdown básico como **negrito**.
    """
    p = doc.add_paragraph()
    # Usa regex para encontrar texto em negrito (**texto**) e texto normal
    parts = re.split(r'(\*\*.*?\*\*)', markdown_text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            # Remove os asteriscos e adiciona o texto como uma "run" em negrito
            p.add_run(part[2:-2]).bold = True
        else:
            p.add_run(part)

def add_markdown_to_doc(doc, markdown_content):
    """
    Analisa um bloco de conteúdo Markdown e o adiciona ao documento .docx,
    lidando com parágrafos, listas e tabelas.
    """
    lines = markdown_content.strip().split('\n')
    in_table = False
    table = None

    for line in lines:
        stripped_line = line.strip()

        if stripped_line.startswith('#### '):
            doc.add_heading(stripped_line.replace('#### ', ''), level=4)
            in_table = False

        # Lida com listas de bullet points
        if stripped_line.startswith(('* ', '- ')):
            text = stripped_line[2:]
            doc.add_paragraph(text, style='List Bullet')
            in_table = False
        # Lida com tabelas
        elif stripped_line.startswith('|') and stripped_line.endswith('|'):
            columns = [cell.strip() for cell in stripped_line.split('|')[1:-1]]
            if not in_table:
                table = doc.add_table(rows=1, cols=len(columns))
                table.style = 'Table Grid'
                hdr_cells = table.rows[0].cells
                for i, col_name in enumerate(columns):
                    hdr_cells[i].text = col_name
                in_table = True
            # Ignora a linha de separador da tabela
            elif not all(c.isspace() or c == '-' for c in columns):
                 row_cells = table.add_row().cells
                 for i, cell_text in enumerate(columns):
                     row_cells[i].text = cell_text
        # Parágrafos normais
        else:
            if stripped_line:
                add_markdown_paragraph(doc, stripped_line)
            in_table = False

def create_final_document(processed_content: dict, output_path: Path):
    """
    Monta o documento .docx final a partir do conteúdo de capítulo completo.
    """
    logger.info("Montando o documento .docx final com formatação rica...")
    doc = docx.Document()
    
    for unit_title, unit_data in processed_content.items():
        doc.add_heading(unit_title, level=1)
        
        if 'theme' in unit_data:
            doc.add_heading("Temáticas da unidade", level=2)
            doc.add_paragraph(unit_data['theme'])
            
        for chapter_title, chapter_data in unit_data.get('chapters', {}).items():
            doc.add_heading(chapter_title, level=2)
            
            # O conteúdo agora é uma string única com todo o capítulo
            if 'content' in chapter_data:
                add_markdown_to_doc(doc, chapter_data['content'])
                
            if 'curiosity' in chapter_data:
                doc.add_heading("Você sabia?", level=3)
                doc.add_paragraph(chapter_data['curiosity'])

            if 'summary' in chapter_data:
                doc.add_heading("Resumindo", level=3)
                add_markdown_to_doc(doc, chapter_data['summary'])
    try:
        doc.save(output_path)
        logger.info(f"Documento final salvo com sucesso em '{output_path}'")
    except PermissionError:
        logger.error(f"ERRO DE PERMISSÃO: Verifique se '{output_path}' não está aberto.")
    except Exception as e:
        logger.error(f"Ocorreu um erro inesperado ao salvar o documento: {e}")