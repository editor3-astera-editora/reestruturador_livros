import pypandoc 
import re 
from pathlib import Path 
from typing import List 
from langchain_core.documents import Document 
from src.utils.logger import logger 

import docx 
from docx.document import Document as DocumentClass 
from docx.oxml.table import CT_Tbl 
from docx.oxml.text.paragraph import CT_P

def remove_section_from_docx(doc: DocumentClass, section_title: str):
    """
    Encontra um título de seção e remove todo o conteúdo a partir dele.
    """
    logger.info(f"Procurando e removendo a seção: '{section_title}'")
    elements_to_remove = []
    found_section = False 

    # Itera sobre todos os elementos do corpo do documento (parágrafos e tabelas)
    for element in doc.element.body:
        if found_section:
            elements_to_remove.append(element)
        elif isinstance(element, CT_P):
            p = docx.text.paragraph.Paragraph(element, doc)
            if section_title.lower() in p.text.lower():
                found_section = True 
                elements_to_remove.append(element)
        
    if found_section:
        for element in elements_to_remove:
            doc.element.body.remove(element)
        logger.info(f"Seção '{section_title}' removida com sucesso.")
    else:
        logger.warning(f"A seção '{section_title}' não foi encontrada no documento.")    

def remove_image_urls_from_docx(doc: DocumentClass):
    """
    Remove parágrafos que contêm URLs de imagens e a linha de fonte correspondente.
    """
    logger.info("Removendo URLs de imagens e fontes do documento...")
    paragraphs_to_remove = []

    for p in doc.paragraphs:
        if re.search(r'https?://\S+', p.text):
            paragraphs_to_remove.append(p)
        elif p.text.strip().lower().startswith('fonte:'):
            paragraphs_to_remove.append(p)

    for p in paragraphs_to_remove:
        parent_element = p._p.getparent()
        if parent_element is not None:
            parent_element.remove(p._p)
            
    logger.info(f"{len(paragraphs_to_remove)} parágrafos contendo URLs/fontes foram removidos.")

def convert_docx_to_markdown(input_path: Path, output_path: Path, intermediate_dir: Path):
    """Converte um arquivo .docx para Markdown usando pandoc."""
    logger.info(f"Convertendo '{input_path.name}' para Markdown...")

    # 1. Carregar o documento 
    doc = docx.Document(input_path)

    # 2. Aplicar as limpezas 
    remove_section_from_docx(doc, "Exercícios resolvidos")
    remove_image_urls_from_docx(doc)

    # 3. Salva um arquivo .docx limpo temporariamente 
    cleaned_docx_path = intermediate_dir / "cleaned_document.docx"
    doc.save(cleaned_docx_path)
    logger.info(f"Documento limpo salvo temporariamente em '{cleaned_docx_path}'")

    # 4. Converter o documento limpo para Markdown
    logger.info(f"Convertendo documento limpo para Markdown...")

    try:
        pypandoc.convert_file(
            source_file=str(cleaned_docx_path),
            to='markdown',
            outputfile=str(output_path),
            extra_args=['--wrap=none'] # evita quebra de linhas indesejadas 
        )
        logger.info(f"Arquivo Markdown salvo em '{output_path}'")
    except Exception as e:
        logger.error(f"Falha ao converter DOCX para Markdown. Certifique-se que o Pandoc está instalado. Erro: {e}")
        raise

def preprocess_markdown_headings(content: str) -> str:
    """
    Corrige o Markdown para usar # e ## ao invés de ** para títulos.
    """
    logger.info("Pré-processando Markdown para corrigir cabeçalhos...")
    lines = content.split('\n')
    processed_lines = []
    for line in lines:
        stripped_line = line.strip()

        # Procura por linhas que são SÓ negrito (ex: **UNIDADE 1**)
        match = re.fullmatch(r'\*\*(.*?)\*\*', stripped_line)
        if match: 
            inner_text = match.group(1).strip()
            # Se for um título de Unidade, usa #
            if inner_text.lower().startswith('unidade'):
                processed_lines.append(f'# {inner_text}')
            # Para outros títulos em negrito, usa ##
            else:
                processed_lines.append(f'## {inner_text}')
        else:
            processed_lines.append(line)
    
    return "\n".join(processed_lines)

def load_and_split_by_structure(content: str, source_name: str) -> List[Document]:
    """
    Divide o conteúdo de uma string Markdown em Documentos, agrupando o conteúdo
    de subtítulos sob seus títulos pais.
    """
    logger.info(f"Estruturando o conteúdo do Markdown de forma hierárquica: '{source_name}'")
    
    # Adiciona uma linha de título final para garantir que a última seção seja capturada
    content += "\n# FIM DO DOCUMENTO"
    lines = content.split('\n')
    
    documents = []
    current_section_content = []
    current_section_title = ""
    current_unit_title = ""

    for line in lines:
        stripped_line = line.strip()
        is_unit_title = stripped_line.startswith('# ')
        is_section_title = stripped_line.startswith('## ')

        # Se a linha for um novo título, salva a seção anterior
        if is_unit_title or is_section_title:
            if current_section_title and current_section_content:
                page_content = "\n".join(current_section_content).strip()
                if page_content:
                    documents.append(Document(
                        page_content=page_content,
                        metadata={'source': source_name, 'title': current_section_title}
                    ))
            
            # Reseta para a nova seção
            current_section_content = []
            if is_unit_title:
                current_unit_title = stripped_line
                current_section_title = current_unit_title # A própria unidade é uma seção
            else: # is_section_title
                # O título completo inclui a unidade para contexto
                current_section_title = f"{current_unit_title}\n{stripped_line}"
        else:
            # É uma linha de conteúdo, adiciona ao bloco atual
            current_section_content.append(line)
    
    # Agora, vamos agrupar o conteúdo.
    # Cria um dicionário de seções, ex: {'# U1\n## S1': 'conteúdo', '# U1\n## S1.1': 'conteúdo'}
    section_dict = {doc.metadata['title']: doc.page_content for doc in documents}
    final_documents = []

    # Itera pelos documentos originais para agrupar o conteúdo
    for doc in documents:
        title = doc.metadata['title']
        content = doc.page_content
        
        # Encontra todos os subtítulos diretos desta seção
        subtitles = [
            other_title for other_title in section_dict 
            if other_title.startswith(title + "\n") and title != other_title
        ]
        
        # Se esta seção tem subtítulos, seu conteúdo final é a junção do seu texto
        # E do texto completo de todos os seus subtítulos.
        if subtitles:
            full_content_list = [content]
            for subtitle in subtitles:
                full_content_list.append(subtitle.split('\n')[-1]) # Adiciona o próprio subtítulo
                full_content_list.append(section_dict[subtitle]) # Adiciona o conteúdo do subtítulo
            
            final_content = "\n\n".join(full_content_list)
            
            # Atualiza o conteúdo ou cria um novo documento
            existing_doc = next((d for d in final_documents if d.metadata['title'] == title), None)
            if existing_doc:
                existing_doc.page_content = final_content
            else:
                final_documents.append(Document(page_content=final_content, metadata=doc.metadata))
        else:
            # Se não tem subtítulos, apenas adiciona se já não foi incluído
            is_subtopic_of_another = any(title.startswith(d.metadata['title'] + "\n") for d in documents if d.metadata['title'] != title)
            if not is_subtopic_of_another:
                final_documents.append(doc)

    logger.info(f"Documento agrupado em {len(final_documents)} seções lógicas principais.")
    return final_documents